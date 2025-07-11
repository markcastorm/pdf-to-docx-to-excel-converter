#!/usr/bin/env python3
"""
Simple DOCX to Excel Table Copier
Just copy tables from DOCX to Excel sheets as-is, no complex processing
"""

import os
import sys
from pathlib import Path
import logging

try:
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install python-docx openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class SimpleDocxToExcel:
    """Copy DOCX tables directly to Excel sheets without overthinking"""
    
    def __init__(self, input_folder: str = "input", output_folder: str = "extracted_data"):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.output_folder.mkdir(exist_ok=True)
        
        # Table names in order (4 per page)
        self.table_names = [
            "Table1_Main_Summary",
            "Table2_Brokerage_Breakdown", 
            "Table3_Institutions_Breakdown",
            "Table4_Financial_Breakdown"
        ]

    def process_all_docx_files(self):
        """Process all DOCX files in input folder"""
        if not self.input_folder.exists():
            logger.error(f"Input folder '{self.input_folder}' does not exist!")
            return
        
        docx_files = list(self.input_folder.glob("*.docx"))
        if not docx_files:
            logger.error(f"No DOCX files found in '{self.input_folder}'")
            return
        
        logger.info(f"Found {len(docx_files)} DOCX files to process...")
        
        for docx_file in docx_files:
            logger.info(f"Processing: {docx_file.name}")
            self.convert_docx_to_excel(docx_file)

    def convert_docx_to_excel(self, docx_path: Path):
        """Convert DOCX tables to Excel sheets"""
        try:
            # Open DOCX document
            doc = Document(docx_path)
            logger.info(f"Found {len(doc.tables)} tables in {docx_path.name}")
            
            # Create Excel workbook
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Create summary sheet first
            self.create_summary_sheet(wb, len(doc.tables), docx_path.name)
            
            # Process each table
            for table_index, table in enumerate(doc.tables):
                # Calculate page and table position
                page_number = (table_index // 4) + 1
                table_position = table_index % 4
                table_name = self.table_names[table_position]
                
                # Create sheet name (shortened to fit Excel limit)
                sheet_name = f"P{page_number}_{table_name[:20]}"  # Limit to ~31 chars total
                
                # Create worksheet
                ws = wb.create_sheet(title=sheet_name)
                
                # Copy table data directly
                self.copy_table_to_sheet(table, ws, docx_path.name, page_number, table_name)
                
                logger.info(f"Copied table {table_index + 1} to sheet '{sheet_name}'")
            
            # Save Excel file
            excel_path = self.output_folder / "multi_table_extraction.xlsx"
            wb.save(excel_path)
            
            logger.info(f"Saved {len(doc.tables)} tables to {excel_path}")
            
            # Print summary
            total_pages = (len(doc.tables) + 3) // 4  # Round up division
            print(f"\n✅ Conversion Complete!")
            print(f"📄 Total Pages: {total_pages}")
            print(f"📊 Total Tables: {len(doc.tables)}")
            print(f"💾 Saved to: {excel_path}")
            
        except Exception as e:
            logger.error(f"Error processing {docx_path.name}: {e}")

    def copy_table_to_sheet(self, table, ws, filename, page_number, table_name):
        """Copy DOCX table to Excel worksheet exactly as-is"""
        # Add metadata
        ws['A1'] = f"File: {filename}"
        ws['A2'] = f"Page: {page_number}"
        ws['A3'] = f"Table: {table_name}"
        ws['A4'] = f"Rows: {len(table.rows)}"
        ws['A5'] = f"Columns: {len(table.columns) if table.rows else 0}"
        
        # Start copying table data from row 7
        start_row = 7
        
        for row_index, table_row in enumerate(table.rows):
            excel_row = start_row + row_index
            
            for col_index, cell in enumerate(table_row.cells):
                excel_col = col_index + 1
                cell_text = cell.text.strip()
                
                # Copy cell content directly
                ws.cell(row=excel_row, column=excel_col, value=cell_text)
        
        # Apply basic formatting
        self.apply_basic_formatting(ws, start_row, len(table.rows))

    def apply_basic_formatting(self, ws, start_row, num_rows):
        """Apply basic formatting to the worksheet"""
        # Format metadata
        for row in ws[1:5]:
            for cell in row:
                if cell.value:
                    cell.font = Font(bold=True)
        
        # Format first row of table (likely headers)
        if num_rows > 0:
            header_row = start_row
            for cell in ws[header_row]:
                if cell.value:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            # Set column width (max 30 chars)
            adjusted_width = min(max_length + 2, 30)
            ws.column_dimensions[column_letter].width = adjusted_width

    def create_summary_sheet(self, wb, total_tables, filename):
        """Create summary sheet"""
        ws = wb.create_sheet(title="Summary", index=0)
        
        ws['A1'] = "DOCX to Excel Direct Copy"
        ws['A1'].font = Font(size=16, bold=True)
        
        ws['A3'] = f"Source File: {filename}"
        ws['A4'] = f"Total Tables: {total_tables}"
        
        total_pages = (total_tables + 3) // 4
        ws['A5'] = f"Total Pages: {total_pages}"
        ws['A6'] = f"Tables per Page: 4"
        
        ws['A8'] = "Sheet Structure:"
        ws['A8'].font = Font(bold=True)
        
        # Show sheet organization
        row = 9
        for page in range(1, min(total_pages + 1, 6)):  # Show first 5 pages
            for i, table_name in enumerate(self.table_names):
                sheet_name = f"P{page}_{table_name[:20]}"
                ws[f'A{row}'] = f"├── {sheet_name}"
                row += 1
        
        if total_pages > 5:
            ws[f'A{row}'] = f"... and {total_pages - 5} more pages"


def main():
    """Main execution function"""
    print("📋 Simple DOCX to Excel Table Copier")
    print("🔄 Copying tables directly without complex processing")
    print("=" * 55)
    
    converter = SimpleDocxToExcel()
    
    if not converter.input_folder.exists():
        print(f"❌ Input folder '{converter.input_folder}' does not exist!")
        print("📁 Please create the folder and place your DOCX files there.")
        return
    
    # Process all DOCX files
    converter.process_all_docx_files()


if __name__ == "__main__":
    main()