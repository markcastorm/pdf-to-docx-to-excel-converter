#!/usr/bin/env python3
"""
PDF to DOCX Conversion and Enhanced Table Extraction
Combines PDF conversion with improved table parsing
"""

import os
import re
import sys
import pandas as pd
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from enum import Enum
import logging

try:
    import pdf2docx
    from docx import Document
    from docx.table import Table as DocxTable
    import pdfplumber
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install pdf2docx python-docx pdfplumber pandas openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ExtractionMethod(Enum):
    """Available extraction methods"""
    DIRECT_PDF = "direct_pdf"
    PDF_TO_DOCX = "pdf_to_docx"
    HYBRID = "hybrid"

class TableType(Enum):
    """Table types in financial reports"""
    TABLE1_MAIN_SUMMARY = "Table1_Main_Summary"
    TABLE2_BROKERAGE_BREAKDOWN = "Table2_Brokerage_Breakdown"
    TABLE3_INSTITUTIONS_BREAKDOWN = "Table3_Institutions_Breakdown"
    TABLE4_FINANCIAL_BREAKDOWN = "Table4_Financial_Breakdown"

@dataclass
class TableData:
    """Extracted table data with metadata"""
    table_type: TableType
    page_number: int
    data: List[Dict[str, Any]]
    categories_found: List[str]
    expected_rows: int
    actual_rows: int
    extraction_method: str
    confidence_score: float

class EnhancedTableExtractor:
    """Multi-method table extractor with PDF to DOCX conversion support"""
    
    def __init__(self, input_folder: str = "input", output_folder: str = "extracted_data"):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.docx_folder = self.output_folder / "converted_docx"
        
        # Create necessary folders
        self.output_folder.mkdir(exist_ok=True)
        self.docx_folder.mkdir(exist_ok=True)
        
        # Table identification patterns
        self.table_patterns = {
            TableType.TABLE1_MAIN_SUMMARY: {
                'japanese_headers': ['総計・自己合計・委託合計', '自己取引計', '委託取引計', '自己委託合計'],
                'english_headers': ['Total, Proprietary', 'Proprietary', 'Brokerage', 'Total'],
                'expected_categories': ['自己取引計', '委託取引計', '自己委託合計'],
                'expected_rows': 9
            },
            TableType.TABLE2_BROKERAGE_BREAKDOWN: {
                'japanese_headers': ['委託内訳', '法人計', '個人計', '海外投資家計', '証券会社'],
                'english_headers': ['Breakdown of Brokerage', 'Institutions', 'Individuals', 'Foreigners', 'Securities Cos'],
                'expected_categories': ['法人計', '個人計', '海外投資家計', '証券会社'],
                'expected_rows': 12
            },
            TableType.TABLE3_INSTITUTIONS_BREAKDOWN: {
                'japanese_headers': ['法人内訳', '投資信託', '事業法人', 'その他法人', '金融機関計'],
                'english_headers': ['Breakdown of Institutions', 'Investment Trusts', 'Business Cos', 'Other Institutions', 'Financial Institutions'],
                'expected_categories': ['投資信託', '事業法人', 'その他法人', '金融機関計'],
                'expected_rows': 12
            },
            TableType.TABLE4_FINANCIAL_BREAKDOWN: {
                'japanese_headers': ['金融機関内訳', '生保・損保', '都銀・地銀等', '信託銀行', 'その他金融機関'],
                'english_headers': ['Breakdown of Financial Institutions', 'Insurance Cos', 'City BKs, Regional BKs', 'Trust Banks', 'Other Financial Institutions'],
                'expected_categories': ['生保・損保', '都銀・地銀等', '信託銀行', 'その他金融機関'],
                'expected_rows': 12
            }
        }
        
        self.transaction_mappings = {
            "売り": "Sales",
            "買い": "Purchases", 
            "合計": "Total"
        }

    def convert_pdf_to_docx(self, pdf_path: Path) -> Path:
        """Convert PDF to DOCX using pdf2docx"""
        docx_path = self.docx_folder / f"{pdf_path.stem}.docx"
        
        if docx_path.exists():
            logger.info(f"DOCX already exists: {docx_path}")
            return docx_path
            
        try:
            logger.info(f"Converting {pdf_path.name} to DOCX...")
            
            # Use pdf2docx for conversion
            from pdf2docx import Converter
            cv = Converter(str(pdf_path))
            cv.convert(str(docx_path), start=0, end=None)
            cv.close()
            
            logger.info(f"Successfully converted to {docx_path}")
            return docx_path
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise

    def extract_from_docx(self, docx_path: Path) -> List[TableData]:
        """Extract tables from DOCX document"""
        try:
            doc = Document(docx_path)
            extracted_tables = []
            
            logger.info(f"Processing DOCX: {docx_path.name}")
            logger.info(f"Found {len(doc.tables)} tables in document")
            
            # Group tables by page (approximate)
            page_number = 1
            tables_per_page = []
            current_page_tables = []
            
            # Process each table in the document
            for table_idx, table in enumerate(doc.tables):
                table_text = self.extract_table_text_from_docx_table(table)
                table_type = self.identify_table_type(table_text)
                
                if table_type:
                    table_data = self.parse_docx_table(table, table_type, page_number)
                    if table_data:
                        extracted_tables.append(table_data)
                        current_page_tables.append(table_data)
                        
                        # Check if we've collected 4 tables (likely a complete page)
                        if len(current_page_tables) >= 4:
                            tables_per_page.append(current_page_tables)
                            current_page_tables = []
                            page_number += 1
            
            # Add any remaining tables
            if current_page_tables:
                tables_per_page.append(current_page_tables)
            
            logger.info(f"Extracted {len(extracted_tables)} tables from {len(tables_per_page)} pages")
            return extracted_tables
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            return []

    def extract_table_text_from_docx_table(self, table: DocxTable) -> List[str]:
        """Extract text content from a DOCX table"""
        text_lines = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                text_lines.append(" ".join(row_text))
        
        return text_lines

    def identify_table_type(self, table_text: List[str]) -> Optional[TableType]:
        """Identify table type from text content"""
        full_text = " ".join(table_text).lower()
        
        for table_type, patterns in self.table_patterns.items():
            # Check Japanese headers
            for header in patterns['japanese_headers']:
                if header in full_text:
                    return table_type
            
            # Check English headers
            for header in patterns['english_headers']:
                if header.lower() in full_text:
                    return table_type
        
        return None

    def parse_docx_table(self, table: DocxTable, table_type: TableType, page_number: int) -> Optional[TableData]:
        """Parse a DOCX table into structured data"""
        try:
            rows_data = []
            categories_found = []
            
            # Extract table data row by row
            for row_idx, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                
                # Skip header rows
                if row_idx < 2:
                    continue
                
                # Look for category and transaction data
                if len(row_cells) >= 6:  # Minimum columns for data rows
                    parsed_row = self.parse_table_row(row_cells, table_type)
                    if parsed_row:
                        rows_data.append(parsed_row)
                        
                        # Track categories
                        category = parsed_row.get('Category')
                        if category and category not in categories_found:
                            categories_found.append(category)
            
            if rows_data:
                return TableData(
                    table_type=table_type,
                    page_number=page_number,
                    data=rows_data,
                    categories_found=categories_found,
                    expected_rows=self.table_patterns[table_type]['expected_rows'],
                    actual_rows=len(rows_data),
                    extraction_method="docx",
                    confidence_score=self.calculate_confidence_score(rows_data, table_type)
                )
            
            return None
            
        except Exception as e:
            logger.error(f"Error parsing DOCX table: {e}")
            return None

    def parse_table_row(self, row_cells: List[str], table_type: TableType) -> Optional[Dict[str, Any]]:
        """Parse a single table row into structured data"""
        if len(row_cells) < 6:
            return None
        
        # Identify category
        category = None
        transaction_type = None
        
        # Look for Japanese category names
        expected_categories = self.table_patterns[table_type]['expected_categories']
        for cat in expected_categories:
            if any(cat in cell for cell in row_cells):
                category = cat
                break
        
        # Look for transaction type
        for jp_trans, en_trans in self.transaction_mappings.items():
            if any(jp_trans in cell or en_trans in cell for cell in row_cells):
                transaction_type = en_trans
                break
        
        if not category or not transaction_type:
            return None
        
        # Extract numerical data
        numbers = []
        for cell in row_cells:
            # Extract numbers (with commas, decimals, and negative signs)
            number_matches = re.findall(r'[▲]?[0-9,]+(?:\.[0-9]+)?', cell)
            numbers.extend(number_matches)
        
        if len(numbers) >= 4:
            return {
                'Category': category,
                'Transaction_Type': transaction_type,
                'Trading Volume Volume': numbers[0] if len(numbers) > 0 else None,
                'Trading Volume Ratio': numbers[1] if len(numbers) > 1 else None,
                'Trading Volume Balance': numbers[2] if len(numbers) > 2 and numbers[2] != '-' else None,
                'Trading Value Value': numbers[3] if len(numbers) > 3 else None,
                'Trading Value Ratio': numbers[4] if len(numbers) > 4 else None,
                'Trading Value Balance': numbers[5] if len(numbers) > 5 and numbers[5] != '-' else None
            }
        
        return None

    def calculate_confidence_score(self, data: List[Dict], table_type: TableType) -> float:
        """Calculate confidence score for extracted data"""
        if not data:
            return 0.0
        
        expected_rows = self.table_patterns[table_type]['expected_rows']
        actual_rows = len(data)
        expected_categories = self.table_patterns[table_type]['expected_categories']
        
        # Score based on row count match
        row_score = min(actual_rows / expected_rows, 1.0)
        
        # Score based on category coverage
        found_categories = set(row['Category'] for row in data if row.get('Category'))
        category_score = len(found_categories) / len(expected_categories)
        
        # Score based on data completeness
        complete_rows = sum(1 for row in data if all(row.get(key) for key in ['Category', 'Transaction_Type', 'Trading Volume Volume', 'Trading Value Value']))
        completeness_score = complete_rows / len(data) if data else 0
        
        return (row_score + category_score + completeness_score) / 3

    def extract_with_fallback(self, pdf_path: Path) -> List[TableData]:
        """Extract using DOCX method with PDF fallback"""
        extracted_tables = []
        
        try:
            # Primary method: PDF to DOCX
            logger.info("Attempting PDF to DOCX conversion method...")
            docx_path = self.convert_pdf_to_docx(pdf_path)
            extracted_tables = self.extract_from_docx(docx_path)
            
            if extracted_tables:
                avg_confidence = sum(t.confidence_score for t in extracted_tables) / len(extracted_tables)
                logger.info(f"DOCX extraction completed with average confidence: {avg_confidence:.2f}")
                
                if avg_confidence > 0.7:  # Good confidence threshold
                    return extracted_tables
                else:
                    logger.warning("Low confidence in DOCX extraction, trying PDF fallback...")
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
        
        # Fallback method: Direct PDF extraction (your original method)
        logger.info("Using direct PDF extraction as fallback...")
        try:
            # You can integrate your original PDF extraction logic here
            # For now, returning empty list
            logger.warning("PDF fallback not implemented in this version")
            return []
            
        except Exception as e:
            logger.error(f"PDF fallback also failed: {e}")
            return []

    def process_all_pdfs(self) -> Dict[str, List[TableData]]:
        """Process all PDFs in input folder"""
        if not self.input_folder.exists():
            logger.error(f"Input folder '{self.input_folder}' does not exist!")
            return {}
        
        pdf_files = list(self.input_folder.glob("*.pdf"))
        if not pdf_files:
            logger.error(f"No PDF files found in '{self.input_folder}'")
            return {}
        
        logger.info(f"Found {len(pdf_files)} PDF files to process...")
        
        all_results = {}
        
        for pdf_file in pdf_files:
            logger.info(f"Processing: {pdf_file.name}")
            try:
                tables_data = self.extract_with_fallback(pdf_file)
                if tables_data:
                    all_results[pdf_file.name] = tables_data
                    logger.info(f"Extracted {len(tables_data)} tables from {pdf_file.name}")
                else:
                    logger.warning(f"No tables extracted from {pdf_file.name}")
                    
            except Exception as e:
                logger.error(f"Error processing {pdf_file.name}: {e}")
                continue
        
        if all_results:
            self.save_to_excel(all_results)
            self.generate_comparison_report(all_results)
            
        return all_results

    def save_to_excel(self, all_results: Dict[str, List[TableData]]) -> None:
        """Save extracted data to Excel with improved formatting"""
        excel_path = self.output_folder / "enhanced_extraction_results.xlsx"
        
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        # Summary data
        summary_data = {
            'total_files': len(all_results),
            'total_tables': sum(len(tables) for tables in all_results.values()),
            'avg_confidence': 0,
            'method_breakdown': {}
        }
        
        all_tables = [table for tables in all_results.values() for table in tables]
        if all_tables:
            summary_data['avg_confidence'] = sum(t.confidence_score for t in all_tables) / len(all_tables)
        
        # Create summary sheet
        self.create_summary_sheet(wb, summary_data, all_results)
        
        # Create data sheets
        for filename, tables_data in all_results.items():
            for table_data in tables_data:
                sheet_name = f"P{table_data.page_number}_{table_data.table_type.value}"
                ws = wb.create_sheet(title=sheet_name)
                self.populate_data_sheet(ws, table_data, filename)
        
        wb.save(excel_path)
        logger.info(f"Enhanced results saved to {excel_path}")

    def create_summary_sheet(self, wb: Workbook, summary_data: Dict, all_results: Dict) -> None:
        """Create comprehensive summary sheet"""
        ws = wb.create_sheet(title="Enhanced_Summary", index=0)
        
        # Headers
        ws['A1'] = "Enhanced Multi-Table Extraction Report"
        ws['A1'].font = Font(size=16, bold=True)
        
        # Summary statistics
        ws['A3'] = "Extraction Summary"
        ws['A3'].font = Font(bold=True)
        
        ws['A4'] = "Total Files Processed:"
        ws['B4'] = summary_data['total_files']
        ws['A5'] = "Total Tables Extracted:"
        ws['B5'] = summary_data['total_tables']
        ws['A6'] = "Average Confidence Score:"
        ws['B6'] = f"{summary_data['avg_confidence']:.2f}"
        
        # Detailed breakdown
        row = 8
        ws[f'A{row}'] = "File-by-File Results:"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        
        for filename, tables in all_results.items():
            ws[f'A{row}'] = filename
            ws[f'B{row}'] = f"{len(tables)} tables"
            avg_conf = sum(t.confidence_score for t in tables) / len(tables) if tables else 0
            ws[f'C{row}'] = f"{avg_conf:.2f} confidence"
            row += 1

    def populate_data_sheet(self, ws, table_data: TableData, filename: str) -> None:
        """Populate individual data sheet with enhanced metadata"""
        # Metadata section
        ws['A1'] = f"File: {filename}"
        ws['A2'] = f"Page: {table_data.page_number}"
        ws['A3'] = f"Table: {table_data.table_type.value}"
        ws['A4'] = f"Extraction Method: {table_data.extraction_method}"
        ws['A5'] = f"Confidence Score: {table_data.confidence_score:.2f}"
        ws['A6'] = f"Expected Rows: {table_data.expected_rows}"
        ws['A7'] = f"Actual Rows: {table_data.actual_rows}"
        ws['A8'] = f"Categories Found: {', '.join(table_data.categories_found)}"
        
        # Data section
        if table_data.data:
            df = pd.DataFrame(table_data.data)
            
            # Add column headers
            start_row = 10
            for col_num, column_name in enumerate(df.columns, 1):
                ws.cell(row=start_row, column=col_num, value=column_name)
            
            # Add data rows
            for row_num, row_data in enumerate(df.values, start_row + 1):
                for col_num, cell_value in enumerate(row_data, 1):
                    ws.cell(row=row_num, column=col_num, value=cell_value)
        
        # Apply formatting
        self.apply_enhanced_formatting(ws)

    def apply_enhanced_formatting(self, ws) -> None:
        """Apply enhanced formatting to worksheet"""
        # Header formatting
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True)
        
        # Metadata formatting
        for row in ws[1:8]:
            for cell in row:
                if cell.value:
                    cell.font = Font(bold=True)
        
        # Column header formatting
        if ws.max_row >= 10:
            for cell in ws[10]:
                if cell.value:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center')

    def generate_comparison_report(self, all_results: Dict[str, List[TableData]]) -> None:
        """Generate comparison report between extraction methods"""
        report_path = self.output_folder / "extraction_quality_report.txt"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("ENHANCED EXTRACTION QUALITY REPORT\n")
            f.write("=" * 50 + "\n\n")
            
            for filename, tables in all_results.items():
                f.write(f"File: {filename}\n")
                f.write("-" * 30 + "\n")
                
                for table in tables:
                    f.write(f"  {table.table_type.value}:\n")
                    f.write(f"    Method: {table.extraction_method}\n")
                    f.write(f"    Confidence: {table.confidence_score:.2f}\n")
                    f.write(f"    Rows: {table.actual_rows}/{table.expected_rows}\n")
                    f.write(f"    Categories: {len(table.categories_found)}\n")
                    f.write("\n")
                
                f.write("\n")
        
        logger.info(f"Quality report saved to {report_path}")


def main():
    """Main execution function"""
    print("🚀 Enhanced Multi-Table Extractor with PDF→DOCX Conversion")
    print("=" * 60)
    
    extractor = EnhancedTableExtractor()
    
    if not extractor.input_folder.exists():
        print(f"❌ Input folder '{extractor.input_folder}' does not exist!")
        print("📁 Please create the folder and place your PDF files there.")
        return
    
    # Process all PDFs
    results = extractor.process_all_pdfs()
    
    if results:
        print(f"\n✅ Enhanced extraction completed!")
        print(f"💾 Results saved to: {extractor.output_folder}")
        print(f"📊 DOCX files in: {extractor.docx_folder}")
        print(f"📈 Quality report generated!")
    else:
        print("❌ No data extracted")


if __name__ == "__main__":
    main()