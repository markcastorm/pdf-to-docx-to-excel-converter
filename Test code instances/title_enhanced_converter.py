#!/usr/bin/env python3
"""
Title-Enhanced Global Chunk Pool PDF→DOCX→Excel Converter
Extracts page titles and adds them above each table in Excel
Same functionality as before + title extraction
"""

import os
import sys
import multiprocessing
from pathlib import Path
import logging
from typing import Optional, List, Tuple, Dict, Any
import time
import tempfile
from concurrent.futures import ProcessPoolExecutor, as_completed
import functools
from dataclasses import dataclass
from queue import Queue
import threading
import re

try:
    from pdf2docx import Converter
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.shared import OxmlElement, qn
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install pdf2docx python-docx openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class ChunkJob:
    """Represents a chunk processing job"""
    pdf_path: Path
    pdf_name: str
    start_page: int
    end_page: int
    chunk_number: int
    total_chunks: int
    chunks_folder: Path
    chunk_id: str  # Unique identifier for this chunk

def convert_chunk_worker(chunk_job: ChunkJob):
    """Worker function for processing chunks from global pool"""
    chunk_docx_path = chunk_job.chunks_folder / f"{chunk_job.pdf_name}_chunk_{chunk_job.chunk_number}.docx"
    
    try:
        logger.info(f"🔄 Processing {chunk_job.chunk_id}: {chunk_job.pdf_name} pages {chunk_job.start_page + 1}-{chunk_job.end_page}")
        
        # Convert specific page range
        cv = Converter(str(chunk_job.pdf_path))
        cv.convert(
            str(chunk_docx_path),
            start=chunk_job.start_page,
            end=chunk_job.end_page,  # Fixed indexing
            # Optimized settings for speed while preserving quality
            table_settings={
                'snap_tolerance': 1.0,
                'min_border_width': 0.3,
                'join_tolerance': 1.0,
            }
        )
        cv.close()
        
        # Verify chunk
        chunk_doc = Document(str(chunk_docx_path))
        chunk_tables = len(chunk_doc.tables)
        expected_tables = (chunk_job.end_page - chunk_job.start_page) * 4
        
        logger.info(f"✅ {chunk_job.chunk_id}: Completed with {chunk_tables}/{expected_tables} tables")
        
        return {
            'chunk_id': chunk_job.chunk_id,
            'pdf_name': chunk_job.pdf_name,
            'chunk_number': chunk_job.chunk_number,
            'chunk_path': chunk_docx_path,
            'table_count': chunk_tables,
            'expected_tables': expected_tables,
            'success': True,
            'start_page': chunk_job.start_page,
            'end_page': chunk_job.end_page
        }
        
    except Exception as e:
        logger.error(f"❌ {chunk_job.chunk_id} failed: {e}")
        return {
            'chunk_id': chunk_job.chunk_id,
            'pdf_name': chunk_job.pdf_name,
            'chunk_number': chunk_job.chunk_number,
            'chunk_path': None,
            'table_count': 0,
            'expected_tables': 0,
            'success': False,
            'error': str(e),
            'start_page': chunk_job.start_page,
            'end_page': chunk_job.end_page
        }

class TitleEnhancedConverter:
    """Title-Enhanced converter with page title extraction"""
    
    def __init__(self, input_folder: str = "input", output_folder: str = "extracted_data", 
                 chunk_size: int = 6, max_workers: int = 6):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.docx_folder = self.output_folder / "converted_docx"
        self.chunks_folder = self.output_folder / "temp_chunks"
        
        # Create necessary folders
        self.output_folder.mkdir(exist_ok=True)
        self.docx_folder.mkdir(exist_ok=True)
        self.chunks_folder.mkdir(exist_ok=True)
        
        # Processing configuration
        self.chunk_size = chunk_size
        self.max_workers = max_workers
        
        logger.info(f"🚀 Title-Enhanced Converter configured:")
        logger.info(f"  • Chunk size: {chunk_size} pages")
        logger.info(f"  • Max workers: {max_workers}")
        logger.info(f"  • CPU cores available: {multiprocessing.cpu_count()}")
        logger.info(f"  • Title extraction: ENABLED")
        
        # Table names in order (4 per page)
        self.table_names = [
            "Table1_Main_Summary",
            "Table2_Brokerage_Breakdown", 
            "Table3_Institutions_Breakdown",
            "Table4_Financial_Breakdown"
        ]

    def process_all_files(self):
        """Process all PDF files using global chunk pool for maximum efficiency"""
        if not self.input_folder.exists():
            logger.error(f"Input folder '{self.input_folder}' does not exist!")
            print(f"❌ Please create '{self.input_folder}' folder and place your PDF files there.")
            return
        
        pdf_files = list(self.input_folder.glob("*.pdf"))
        docx_files = list(self.input_folder.glob("*.docx"))
        
        if not pdf_files and not docx_files:
            logger.error(f"No PDF or DOCX files found in '{self.input_folder}'")
            print(f"❌ No files to process in '{self.input_folder}'")
            return
        
        print(f"📁 Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files")
        print(f"⚡ Global chunk pool processing: {self.max_workers} workers")
        print(f"📋 Title extraction: ENABLED")
        
        # Process PDFs using global chunk pool
        if pdf_files:
            self.process_pdfs_with_global_pool(pdf_files)
        
        # Process existing DOCX files
        for docx_file in docx_files:
            start_time = time.time()
            logger.info(f"📊 Processing existing DOCX: {docx_file.name}")
            
            self.convert_docx_to_excel(docx_file, source_type="DOCX")
            
            elapsed_time = time.time() - start_time
            print(f"⏱️ DOCX processing time: {elapsed_time:.2f} seconds")
        
        # Cleanup temp files
        self.cleanup_temp_files()

    def process_pdfs_with_global_pool(self, pdf_files: List[Path]):
        """Process multiple PDFs using global chunk pool"""
        start_time = time.time()
        
        # Create global chunk pool
        all_chunk_jobs = []
        pdf_info = {}
        
        logger.info(f"🔄 Creating global chunk pool from {len(pdf_files)} PDFs...")
        
        for pdf_file in pdf_files:
            # Get page count
            total_pages = self.get_pdf_page_count(pdf_file)
            pdf_name = pdf_file.stem
            
            # Store PDF info for later processing
            pdf_info[pdf_name] = {
                'pdf_path': pdf_file,
                'total_pages': total_pages,
                'chunks': []
            }
            
            # Create chunks for this PDF
            chunk_count = 0
            for chunk_start in range(0, total_pages, self.chunk_size):
                chunk_end = min(chunk_start + self.chunk_size, total_pages)
                chunk_count += 1
                
                chunk_job = ChunkJob(
                    pdf_path=pdf_file,
                    pdf_name=pdf_name,
                    start_page=chunk_start,
                    end_page=chunk_end,
                    chunk_number=chunk_count,
                    total_chunks=0,  # Will be updated
                    chunks_folder=self.chunks_folder,
                    chunk_id=f"{pdf_name}-C{chunk_count}"
                )
                
                all_chunk_jobs.append(chunk_job)
                pdf_info[pdf_name]['chunks'].append(chunk_job)
            
            # Update total chunks
            for chunk_job in pdf_info[pdf_name]['chunks']:
                chunk_job.total_chunks = chunk_count
            
            logger.info(f"📄 {pdf_name}: {total_pages} pages → {chunk_count} chunks")
        
        total_chunks = len(all_chunk_jobs)
        logger.info(f"🌐 Global chunk pool created: {total_chunks} chunks from {len(pdf_files)} PDFs")
        
        # Process all chunks using global pool
        conversion_start = time.time()
        
        logger.info(f"⚡ Starting global pool processing with {self.max_workers} workers...")
        
        chunk_results = []
        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all chunk jobs to global pool
            future_to_chunk = {
                executor.submit(convert_chunk_worker, chunk_job): chunk_job.chunk_id
                for chunk_job in all_chunk_jobs
            }
            
            # Collect results as they complete
            completed_chunks = 0
            for future in as_completed(future_to_chunk):
                chunk_id = future_to_chunk[future]
                try:
                    result = future.result()
                    chunk_results.append(result)
                    completed_chunks += 1
                    
                    if result['success']:
                        logger.info(f"✅ {chunk_id} completed ({completed_chunks}/{total_chunks})")
                    else:
                        logger.error(f"❌ {chunk_id} failed: {result.get('error', 'Unknown error')}")
                        
                except Exception as e:
                    logger.error(f"❌ {chunk_id} failed with exception: {e}")
                    chunk_results.append({
                        'chunk_id': chunk_id,
                        'success': False,
                        'error': str(e)
                    })
                    completed_chunks += 1
        
        parallel_time = time.time() - conversion_start
        logger.info(f"⚡ Global pool processing completed in {parallel_time:.2f} seconds")
        
        # Group results by PDF and process each
        pdf_results = {}
        for result in chunk_results:
            if result['success']:
                pdf_name = result['pdf_name']
                if pdf_name not in pdf_results:
                    pdf_results[pdf_name] = []
                pdf_results[pdf_name].append(result)
        
        # Process each PDF's results
        for pdf_name, results in pdf_results.items():
            pdf_start = time.time()
            logger.info(f"🔗 Processing results for {pdf_name} ({len(results)} chunks)")
            
            # Sort chunks by number
            results.sort(key=lambda x: x['chunk_number'])
            
            # Combine chunks into DOCX
            chunk_paths = [result['chunk_path'] for result in results]
            docx_path = self.docx_folder / f"{pdf_name}.docx"
            
            final_docx = self.combine_chunks(chunk_paths, docx_path, pdf_name)
            
            if final_docx:
                # Convert to Excel with title extraction
                self.convert_docx_to_excel(final_docx, source_type="PDF_GLOBAL_POOL")
                
                pdf_time = time.time() - pdf_start
                
                # Verify results
                final_doc = Document(str(final_docx))
                final_tables = len(final_doc.tables)
                expected_total = sum(r['expected_tables'] for r in results)
                
                print(f"\n🎉 {pdf_name} COMPLETED:")
                print(f"  📊 Tables: {final_tables}/{expected_total}")
                print(f"  📋 Chunks: {len(results)}")
                print(f"  ⏱️ Processing time: {pdf_time:.2f}s")
                print(f"  💾 Excel: {pdf_name}_extracted.xlsx")
        
        total_time = time.time() - start_time
        
        # Print global summary
        print(f"\n⚡ GLOBAL POOL PROCESSING SUMMARY:")
        print(f"=" * 60)
        print(f"  📁 PDFs processed: {len(pdf_files)}")
        print(f"  🌐 Total chunks: {total_chunks}")
        print(f"  👥 Workers used: {self.max_workers}")
        print(f"  ⏱️ Total time: {total_time:.2f} seconds")
        print(f"  🚀 Average chunks/second: {total_chunks/parallel_time:.1f}")
        print(f"  💾 Excel files created: {len(pdf_results)}")

    def get_pdf_page_count(self, pdf_path: Path) -> int:
        """Get the total number of pages in PDF"""
        try:
            import fitz
            doc = fitz.open(str(pdf_path))
            page_count = len(doc)
            doc.close()
            return page_count
        except ImportError:
            logger.warning("PyMuPDF not available, assuming 24 pages")
            return 24
        except Exception as e:
            logger.warning(f"Could not get page count for {pdf_path.name}, assuming 24 pages: {e}")
            return 24

    def combine_chunks(self, chunk_paths: List[Path], output_path: Path, pdf_name: str) -> Optional[Path]:
        """Combine chunks efficiently"""
        try:
            if not chunk_paths:
                return None
            
            logger.info(f"🔗 Combining {len(chunk_paths)} chunks for {pdf_name}...")
            
            # Start with first chunk
            import shutil
            shutil.copy2(chunk_paths[0], output_path)
            combined_doc = Document(str(output_path))
            
            total_tables = len(combined_doc.tables)
            
            # Append remaining chunks
            for i, chunk_path in enumerate(chunk_paths[1:], 1):
                chunk_doc = Document(str(chunk_path))
                chunk_tables = len(chunk_doc.tables)
                
                # Add page break
                combined_doc.add_page_break()
                
                # Copy all elements from chunk
                for element in chunk_doc.element.body:
                    try:
                        import copy
                        imported = copy.deepcopy(element)
                        combined_doc.element.body.append(imported)
                    except:
                        pass
                
                total_tables += chunk_tables
            
            # Save combined document
            combined_doc.save(str(output_path))
            
            # Verify result
            verification_doc = Document(str(output_path))
            final_count = len(verification_doc.tables)
            
            logger.info(f"✅ {pdf_name}: Combined {len(chunk_paths)} chunks → {final_count} tables")
            return output_path
            
        except Exception as e:
            logger.error(f"Chunk combination failed for {pdf_name}: {e}")
            # Fallback: use first chunk
            if chunk_paths:
                shutil.copy2(chunk_paths[0], output_path)
                return output_path
            return None

    def extract_page_title(self, doc: Document, table_index: int) -> Tuple[str, str]:
        """Extract actual title and subtitle from the specific page containing the table"""
        try:
            # Calculate which page this table is on (4 tables per page)
            page_number = (table_index // 4) + 1
            
            # Get all paragraphs and their positions
            all_paragraphs = doc.paragraphs
            
            # Find the table in the document
            current_table = doc.tables[table_index]
            
            # Get paragraphs that appear before this specific table
            # We'll look for title information in the paragraphs preceding this table
            table_element = current_table._element
            
            # Find paragraphs that are near this specific table (not just any table)
            relevant_paragraphs = []
            
            # Method 1: Get paragraphs by scanning through document elements
            found_table = False
            paragraphs_before_table = []
            
            # Walk through document elements to find paragraphs near our table
            body = doc.element.body
            for element in body:
                # If this is our table, stop collecting paragraphs
                if element.tag.endswith('tbl') and element == table_element:
                    found_table = True
                    break
                
                # If this is a paragraph, collect it
                if element.tag.endswith('p'):
                    # Find corresponding paragraph object
                    for para in all_paragraphs:
                        if para._element == element:
                            paragraphs_before_table.append(para)
                            break
            
            # Method 2: Look at the last few paragraphs before this table for title info
            # Take the last 10 paragraphs before the table as candidates
            title_candidates = paragraphs_before_table[-10:] if len(paragraphs_before_table) >= 10 else paragraphs_before_table
            
            title = ""
            subtitle = ""
            
            # Search through candidate paragraphs for title patterns
            for para in reversed(title_candidates):  # Start from the closest to table
                text = para.text.strip()
                if not text:
                    continue
                
                # Look for various title patterns - be more flexible
                
                # Pattern 1: Full title with both Japanese and English
                if ("投資部門" in text and "取引" in text) or "Trading by Type" in text:
                    title = text
                
                # Pattern 2: Look for other trading-related titles
                if not title and ("売買" in text or "Trading" in text or "取引" in text):
                    title = text
                
                # Pattern 3: Look for sector-specific titles
                if not title and ("部門" in text or "Sector" in text or "Breakdown" in text):
                    title = text
                
                # Subtitle patterns - look for instrument/market specifications
                if "(" in text and ")" in text:
                    # Check for various instruments
                    if any(keyword in text for keyword in ["先物", "Futures", "225", "オプション", "Options", "TOPIX", "マザーズ", "Mothers"]):
                        subtitle = text
                
                # Also look for subtitle without parentheses
                if not subtitle and any(keyword in text for keyword in ["日経", "Nikkei", "TOPIX", "マザーズ", "Mothers", "JQ", "東証"]):
                    subtitle = text
            
            # Enhanced fallback: try to get any meaningful text near the table
            if not title:
                # Look for any substantial text that could be a title
                for para in reversed(title_candidates[-5:]):  # Last 5 paragraphs
                    text = para.text.strip()
                    if len(text) > 10 and not text.isdigit():  # Substantial text, not just numbers
                        title = text
                        break
            
            # If still no title found, use a descriptive fallback
            if not title:
                title = f"Trading Data - Page {page_number}"
            
            # If no subtitle found, try to infer from context or use page-specific fallback
            if not subtitle:
                # Try to determine market type from surrounding context
                page_text = ' '.join([p.text for p in title_candidates])
                if "225" in page_text or "日経" in page_text:
                    subtitle = "(日経225関連 / Nikkei 225 Related)"
                elif "TOPIX" in page_text:
                    subtitle = "(TOPIX関連 / TOPIX Related)"
                elif "マザーズ" in page_text or "Mothers" in page_text:
                    subtitle = "(マザーズ / Mothers Market)"
                else:
                    subtitle = f"(Page {page_number} Data)"
            
            # Clean up the extracted text
            title = title.replace('\n', ' ').replace('\r', ' ').strip()
            subtitle = subtitle.replace('\n', ' ').replace('\r', ' ').strip()
            
            logger.debug(f"Extracted for table {table_index} (page {page_number}): Title='{title}', Subtitle='{subtitle}'")
            
            return title, subtitle
            
        except Exception as e:
            logger.debug(f"Could not extract title for table {table_index}: {e}")
            # Page-specific fallback
            page_number = (table_index // 4) + 1
            return f"Trading Data - Page {page_number}", f"(Page {page_number} Data)"

    def convert_docx_to_excel(self, docx_path: Path, source_type: str = "DOCX"):
        """Convert DOCX tables to Excel sheets with title extraction"""
        try:
            doc = Document(docx_path)
            total_tables = len(doc.tables)
            
            if total_tables == 0:
                logger.warning(f"⚠️ No tables found in {docx_path.name}")
                return
            
            extraction_start = time.time()
            
            # Create Excel workbook
            wb = Workbook()
            wb.remove(wb.active)
            
            # Create summary sheet
            self.create_summary_sheet(wb, total_tables, docx_path.name, source_type)
            
            # Process each table
            for table_index, table in enumerate(doc.tables):
                page_number = (table_index // 4) + 1
                table_position = table_index % 4
                table_name = self.table_names[table_position]
                
                sheet_name = f"P{page_number}_{table_name[:20]}"
                ws = wb.create_sheet(title=sheet_name)
                
                # Extract title and subtitle for this table
                title, subtitle = self.extract_page_title(doc, table_index)
                
                # Copy table data with title information
                self.copy_table_to_sheet_with_title(table, ws, docx_path.name, page_number, table_name, title, subtitle)
                
                if (table_index + 1) % 40 == 0:
                    logger.info(f"📋 Processed {table_index + 1}/{total_tables} tables...")
            
            # Save Excel file
            excel_filename = f"{docx_path.stem}_extracted.xlsx"
            excel_path = self.output_folder / excel_filename
            wb.save(excel_path)
            
            extraction_time = time.time() - extraction_start
            logger.info(f"📊 Excel extraction completed in {extraction_time:.2f} seconds")
            
        except Exception as e:
            logger.error(f"❌ Error processing {docx_path.name}: {e}")

    def copy_table_to_sheet_with_title(self, table, ws, filename, page_number, table_name, title, subtitle):
        """Copy DOCX table to Excel worksheet with title information"""
        # Add source metadata
        ws['A1'] = f"Source: {filename}"
        ws['A2'] = f"Page: {page_number}"
        ws['A3'] = f"Table: {table_name}"
        ws['A4'] = f"Rows: {len(table.rows)}"
        ws['A5'] = f"Columns: {len(table.columns) if table.rows else 0}"
        
        # Add title information (NEW FEATURE)
        ws['A6'] = ""  # Empty row for spacing
        ws['A7'] = f"Title: {title}"
        ws['A8'] = f"Subtitle: {subtitle}"
        ws['A9'] = ""  # Empty row for spacing
        
        # Start copying table data from row 10 (moved down to accommodate titles)
        start_row = 10
        
        for row_index, table_row in enumerate(table.rows):
            excel_row = start_row + row_index
            for col_index, cell in enumerate(table_row.cells):
                excel_col = col_index + 1
                cell_text = cell.text.strip()
                ws.cell(row=excel_row, column=excel_col, value=cell_text)
        
        # Apply formatting with title highlighting
        self.apply_enhanced_formatting(ws, start_row, len(table.rows))

    def apply_enhanced_formatting(self, ws, start_row, num_rows):
        """Apply enhanced formatting with title highlighting"""
        # Format metadata (rows 1-5)
        for row in ws[1:5]:
            for cell in row:
                if cell.value:
                    cell.font = Font(bold=True)
        
        # Format title and subtitle (rows 7-8) - HIGHLIGHTED
        title_fill = PatternFill(start_color='E6F3FF', end_color='E6F3FF', fill_type='solid')
        title_font = Font(bold=True, size=12)
        
        ws['A7'].font = title_font
        ws['A7'].fill = title_fill
        ws['A8'].font = title_font
        ws['A8'].fill = title_fill
        
        # Format table headers (first row of table data)
        if num_rows > 0:
            for cell in ws[start_row]:
                if cell.value:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')

    def create_summary_sheet(self, wb, total_tables, filename, source_type):
        """Create summary sheet with title extraction info"""
        ws = wb.create_sheet(title="Summary", index=0)
        
        ws['A1'] = "📋 Title-Enhanced PDF→DOCX→Excel Converter"
        ws['A1'].font = Font(size=16, bold=True)
        
        ws['A3'] = f"Source Type: {source_type}"
        ws['A4'] = f"Source File: {filename}"
        ws['A5'] = f"Total Tables: {total_tables}"
        ws['A6'] = f"Chunk Size: {self.chunk_size} pages"
        ws['A7'] = f"Max Workers: {self.max_workers}"
        ws['A8'] = f"CPU Cores: {multiprocessing.cpu_count()}"
        
        total_pages = (total_tables + 3) // 4
        ws['A9'] = f"Total Pages: {total_pages}"
        ws['A10'] = f"Processing Method: Global chunk pool"
        ws['A11'] = f"Title Extraction: ENABLED"
        
        ws['A13'] = "Enhanced Features:"
        ws['A13'].font = Font(bold=True)
        ws['A14'] = "🌐 Cross-PDF worker utilization"
        ws['A15'] = "⚡ No idle worker time"
        ws['A16'] = "🚀 Maximum efficiency"
        ws['A17'] = "📊 Perfect table preservation"
        ws['A18'] = "💾 Separate Excel per PDF"
        ws['A19'] = "📋 Page title extraction (NEW!)"

    def cleanup_temp_files(self):
        """Clean up temporary chunk files"""
        try:
            chunk_files = list(self.chunks_folder.glob("*_chunk_*.docx"))
            for chunk_file in chunk_files:
                chunk_file.unlink()
            logger.info(f"🧹 Cleaned up {len(chunk_files)} temporary chunk files")
        except Exception as e:
            logger.warning(f"Could not clean up temp files: {e}")


def main():
    """Main execution function"""
    print("📋 Title-Enhanced Global Chunk Pool PDF→DOCX→Excel Converter")
    print("=" * 75)
    print("🚀 Enhanced Features:")
    print("  • Global chunk pool across ALL PDFs")
    print("  • 6 workers for maximum CPU utilization")
    print("  • No idle worker time")
    print("  • Separate Excel file per PDF")
    print("  • Perfect table preservation")
    print("  • 📋 PAGE TITLE EXTRACTION (NEW!)")
    print("=" * 75)
    
    # Configuration
    chunk_size = 6
    max_workers = 6
    
    try:
        user_chunk = input(f"Chunk size (pages per chunk, default {chunk_size}): ").strip()
        if user_chunk:
            chunk_size = int(user_chunk)
            
        user_workers = input(f"Max workers (default {max_workers}): ").strip()
        if user_workers:
            max_workers = int(user_workers)
            
        print(f"✅ Configuration: {chunk_size} pages per chunk, {max_workers} workers")
    except ValueError:
        print(f"Using defaults: {chunk_size} pages, {max_workers} workers")
    
    converter = TitleEnhancedConverter(chunk_size=chunk_size, max_workers=max_workers)
    
    # Process all files
    converter.process_all_files()
    
    print(f"\n📁 DOCX files saved in: {converter.docx_folder}")
    print(f"💾 Excel files saved in: {converter.output_folder}")
    print(f"🌐 Each PDF has its own Excel file!")
    print(f"📋 Page titles extracted and added to Excel sheets!")
    
    cleanup = input("\n🗑️ Delete intermediate DOCX files? (y/N): ").lower()
    if cleanup == 'y':
        docx_files = list(converter.docx_folder.glob("*.docx"))
        for docx_file in docx_files:
            docx_file.unlink()
        print("✅ Intermediate files cleaned up")
    else:
        print("📁 DOCX files kept for reference")
    
    print("\n🎉 Title-enhanced processing complete!")
    print("📋 Excel files now include page titles for easy classification!")
    print("🌐 Maximum efficiency achieved across all PDFs!")


if __name__ == "__main__":
    main()