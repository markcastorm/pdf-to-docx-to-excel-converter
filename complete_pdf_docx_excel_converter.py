#!/usr/bin/env python3
"""
Complete PDF to DOCX to Excel Converter Pipeline
1. Converts PDF files to DOCX (preserves table structure)
2. Extracts all tables from DOCX to organized Excel sheets
3. No complex parsing - direct table copying for perfect results
"""

import os
import sys
from pathlib import Path
import logging
from typing import Optional

try:
    from pdf2docx import Converter
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install pdf2docx python-docx openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class CompletePdfToExcelConverter:
    """Complete pipeline: PDF → DOCX → Excel with perfect table extraction"""
    
    def __init__(self, input_folder: str = "input", output_folder: str = "extracted_data"):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.docx_folder = self.output_folder / "converted_docx"
        
        # Create necessary folders
        self.output_folder.mkdir(exist_ok=True)
        self.docx_folder.mkdir(exist_ok=True)
        
        # Table names in order (4 per page)
        self.table_names = [
            "Table1_Main_Summary",
            "Table2_Brokerage_Breakdown", 
            "Table3_Institutions_Breakdown",
            "Table4_Financial_Breakdown"
        ]

    def process_all_files(self):
        """Process all PDF files through the complete pipeline"""
        if not self.input_folder.exists():
            logger.error(f"Input folder '{self.input_folder}' does not exist!")
            print(f"❌ Please create '{self.input_folder}' folder and place your PDF files there.")
            return
        
        pdf_files = list(self.input_folder.glob("*.pdf"))
        docx_files = list(self.input_folder.glob("*.docx"))
        
        if not pdf_files and not docx_files:
            logger.error(f"No PDF or DOCX files found in '{self.input_folder}'")
            print(f"❌ No files to process in '{self.input_folder}'")
            return
        
        print(f"📁 Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files")
        
        # Process PDF files first (convert to DOCX)
        for pdf_file in pdf_files:
            logger.info(f"Processing PDF: {pdf_file.name}")
            docx_path = self.convert_pdf_to_docx(pdf_file)
            if docx_path:
                self.convert_docx_to_excel(docx_path, source_type="PDF")
        
        # Process existing DOCX files
        for docx_file in docx_files:
            logger.info(f"Processing DOCX: {docx_file.name}")
            self.convert_docx_to_excel(docx_file, source_type="DOCX")

    def convert_pdf_to_docx(self, pdf_path: Path) -> Optional[Path]:
        """Convert PDF to DOCX using pdf2docx"""
        docx_path = self.docx_folder / f"{pdf_path.stem}.docx"
        
        if docx_path.exists():
            logger.info(f"DOCX already exists: {docx_path.name}")
            return docx_path
            
        try:
            logger.info(f"🔄 Converting {pdf_path.name} to DOCX...")
            
            # Use pdf2docx for conversion
            cv = Converter(str(pdf_path))
            cv.convert(str(docx_path), start=0, end=None)
            cv.close()
            
            logger.info(f"✅ Successfully converted to {docx_path.name}")
            return docx_path
            
        except Exception as e:
            logger.error(f"❌ PDF to DOCX conversion failed for {pdf_path.name}: {e}")
            return None

    def convert_docx_to_excel(self, docx_path: Path, source_type: str = "DOCX"):
        """Convert DOCX tables to Excel sheets (the proven perfect method)"""
        try:
            # Open DOCX document
            doc = Document(docx_path)
            logger.info(f"📊 Found {len(doc.tables)} tables in {docx_path.name}")
            
            if len(doc.tables) == 0:
                logger.warning(f"⚠️ No tables found in {docx_path.name}")
                return
            
            # Create Excel workbook
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Create summary sheet first
            self.create_summary_sheet(wb, len(doc.tables), docx_path.name, source_type)
            
            # Process each table
            for table_index, table in enumerate(doc.tables):
                # Calculate page and table position
                page_number = (table_index // 4) + 1
                table_position = table_index % 4
                table_name = self.table_names[table_position]
                
                # Create sheet name (shortened to fit Excel limit)
                sheet_name = f"P{page_number}_{table_name[:20]}"  # Limit to ~31 chars total
                
                # Create worksheet
                ws = wb.create_sheet(title=sheet_name)
                
                # Copy table data directly (the perfect method!)
                self.copy_table_to_sheet(table, ws, docx_path.name, page_number, table_name)
                
                if (table_index + 1) % 20 == 0:  # Progress update every 20 tables
                    logger.info(f"📋 Processed {table_index + 1}/{len(doc.tables)} tables...")
            
            # Save Excel file
            excel_filename = f"{docx_path.stem}_extracted.xlsx"
            excel_path = self.output_folder / excel_filename
            wb.save(excel_path)
            
            # Print success summary
            total_pages = (len(doc.tables) + 3) // 4  # Round up division
            print(f"\n🎉 SUCCESS! Conversion Complete for {docx_path.name}")
            print(f"📄 Source: {source_type}")
            print(f"📊 Tables Extracted: {len(doc.tables)}")
            print(f"📋 Pages Processed: {total_pages}")
            print(f"💾 Excel File: {excel_path}")
            print("=" * 60)
            
            logger.info(f"✅ Saved {len(doc.tables)} tables to {excel_path}")
            
        except Exception as e:
            logger.error(f"❌ Error processing {docx_path.name}: {e}")

    def copy_table_to_sheet(self, table, ws, filename, page_number, table_name):
        """Copy DOCX table to Excel worksheet exactly as-is (proven perfect method)"""
        # Add metadata
        ws['A1'] = f"Source: {filename}"
        ws['A2'] = f"Page: {page_number}"
        ws['A3'] = f"Table: {table_name}"
        ws['A4'] = f"Rows: {len(table.rows)}"
        ws['A5'] = f"Columns: {len(table.columns) if table.rows else 0}"
        
        # Start copying table data from row 7
        start_row = 7
        
        for row_index, table_row in enumerate(table.rows):
            excel_row = start_row + row_index
            
            for col_index, cell in enumerate(table_row.cells):
                excel_col = col_index + 1
                cell_text = cell.text.strip()
                
                # Copy cell content directly - no processing, no filtering
                ws.cell(row=excel_row, column=excel_col, value=cell_text)
        
        # Apply basic formatting
        self.apply_basic_formatting(ws, start_row, len(table.rows))

    def apply_basic_formatting(self, ws, start_row, num_rows):
        """Apply basic formatting to the worksheet"""
        # Format metadata
        for row in ws[1:5]:
            for cell in row:
                if cell.value:
                    cell.font = Font(bold=True)
        
        # Format first row of table (likely headers)
        if num_rows > 0:
            header_row = start_row
            for cell in ws[header_row]:
                if cell.value:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            # Set column width (max 30 chars)
            adjusted_width = min(max_length + 2, 30)
            ws.column_dimensions[column_letter].width = adjusted_width

    def create_summary_sheet(self, wb, total_tables, filename, source_type):
        """Create summary sheet with conversion info"""
        ws = wb.create_sheet(title="Summary", index=0)
        
        ws['A1'] = "PDF → DOCX → Excel Conversion Results"
        ws['A1'].font = Font(size=16, bold=True)
        
        ws['A3'] = f"Source Type: {source_type}"
        ws['A4'] = f"Source File: {filename}"
        ws['A5'] = f"Total Tables: {total_tables}"
        
        total_pages = (total_tables + 3) // 4
        ws['A6'] = f"Total Pages: {total_pages}"
        ws['A7'] = f"Tables per Page: 4"
        ws['A8'] = f"Extraction Method: Direct table copy (no parsing)"
        
        ws['A10'] = "Conversion Pipeline:"
        ws['A10'].font = Font(bold=True)
        
        if source_type == "PDF":
            ws['A11'] = "1. PDF → DOCX (pdf2docx conversion)"
            ws['A12'] = "2. DOCX → Excel (direct table copy)"
        else:
            ws['A11'] = "1. DOCX → Excel (direct table copy)"
        
        ws['A14'] = "Sheet Structure:"
        ws['A14'].font = Font(bold=True)
        
        # Show sheet organization (first few pages)
        row = 15
        for page in range(1, min(total_pages + 1, 6)):  # Show first 5 pages
            for i, table_name in enumerate(self.table_names):
                sheet_name = f"P{page}_{table_name[:20]}"
                ws[f'A{row}'] = f"├── {sheet_name}"
                row += 1
        
        if total_pages > 5:
            ws[f'A{row}'] = f"... and {total_pages - 5} more pages"

    def cleanup_intermediate_files(self, keep_docx: bool = True):
        """Clean up intermediate DOCX files if desired"""
        if not keep_docx:
            docx_files = list(self.docx_folder.glob("*.docx"))
            for docx_file in docx_files:
                docx_file.unlink()
                logger.info(f"Cleaned up {docx_file.name}")


def main():
    """Main execution function"""
    print("🚀 Complete PDF → DOCX → Excel Converter Pipeline")
    print("=" * 60)
    print("📋 Features:")
    print("  • Converts PDF to DOCX (preserves table structure)")
    print("  • Extracts all tables to organized Excel sheets") 
    print("  • No complex parsing - direct table copying")
    print("  • Handles both PDF and DOCX input files")
    print("  • Proven perfect results!")
    print("=" * 60)
    
    converter = CompletePdfToExcelConverter()
    
    # Process all files
    converter.process_all_files()
    
    # Ask about cleanup
    print(f"\n📁 DOCX files saved in: {converter.docx_folder}")
    print(f"💾 Excel files saved in: {converter.output_folder}")
    
    cleanup = input("\n🗑️ Delete intermediate DOCX files? (y/N): ").lower()
    if cleanup == 'y':
        converter.cleanup_intermediate_files(keep_docx=False)
        print("✅ Intermediate files cleaned up")
    else:
        print("📁 DOCX files kept for reference")
    
    print("\n🎉 Pipeline complete! Check your Excel files!")


if __name__ == "__main__":
    main()