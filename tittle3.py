#!/usr/bin/env python3
"""
Title-Enhanced Selective Page PDF→DOCX→Excel Converter
Pre-scans PDFs to find pages with specific subtitles and converts only those pages.
Maintains original folder structure: input/ -> extracted_data/
"""

import os
import sys
import multiprocessing
from pathlib import Path
import logging
from typing import Optional, List, Tuple, Dict
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass
import re

try:
    import fitz  # PyMuPDF
    from pdf2docx import Converter
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install PyMuPDF pdf2docx python-docx openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


@dataclass
class PageJob:
    """Represents a single-page processing job"""
    pdf_path: Path
    pdf_name: str
    page_number: int  # 0-indexed page number
    subtitle: str     # The subtitle found on this page
    job_id: str       # Unique identifier for this job
    chunks_folder: Path
    table_titles: List[str] = None


def convert_page_worker(job: PageJob):
    """Worker function for converting a single, pre-identified relevant page."""
    page_docx_path = job.chunks_folder / f"{job.pdf_name}_page_{job.page_number}.docx"
    
    try:
        logger.info(f"🔄 Converting {job.job_id}: Page {job.page_number + 1} ('{job.subtitle}')")
        
        cv = Converter(str(job.pdf_path))
        cv.convert(
            str(page_docx_path),
            start=job.page_number,
            end=job.page_number + 1,
            table_settings={'snap_tolerance': 1.0, 'min_border_width': 0.3, 'join_tolerance': 1.0}
        )
        cv.close()
        
        chunk_doc = Document(str(page_docx_path))
        table_count = len(chunk_doc.tables)
        
        logger.info(f"✅ {job.job_id}: Completed with {table_count} tables.")
        
        return {
            'job_id': job.job_id,
            'pdf_name': job.pdf_name,
            'page_number': job.page_number,
            'page_path': page_docx_path,
            'subtitle': job.subtitle,
            'table_count': table_count,
            'table_titles': job.table_titles,
            'success': True
        }
        
    except Exception as e:
        logger.error(f"❌ {job.job_id} failed: {e}")
        return {
            'job_id': job.job_id,
            'pdf_name': job.pdf_name,
            'page_number': job.page_number,
            'success': False,
            'error': str(e)
        }


class TitleEnhancedConverter:
    """Selectively converts PDF pages based on subtitles."""
    
    def __init__(self, input_folder: str = "input", output_folder: str = "extracted_data", max_workers: int = 6):
        # ORIGINAL FOLDER STRUCTURE MAINTAINED
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.docx_folder = self.output_folder / "converted_docx"
        self.chunks_folder = self.output_folder / "temp_chunks"
        
        # Create necessary folders
        self.output_folder.mkdir(exist_ok=True)
        self.docx_folder.mkdir(exist_ok=True)
        self.chunks_folder.mkdir(exist_ok=True)
        
        self.max_workers = max_workers
        
        # Enhanced subtitle detection - captures full subtitle lines
        self.target_subtitles = [
            "長期国債先物（現金決済型ミニ）",
            "JGB(10-year) Futures",
            "mini-10-year JGB Futures (Cash-Settled)",
            "mini-20-year JGB Futures",
            "3-Month TONA Futures"
        ]
        
        # Define keywords that identify an "Options" page to be excluded
        self.primary_exclusion_keywords = ["Options on", "オプション"]

        # Sort targets by length (longest first) to prioritize more specific matches
        self.target_subtitles.sort(key=len, reverse=True)

        # Enhanced table title detection
        self.table_section_titles = [
            "総計・自己合計・委託合計 Total, Proprietary & Brokerage",
            "委託内訳 Breakdown of Brokerage",
            "法人内訳 Breakdown of Institutions",
            "金融機関内訳 Breakdown of Financial Institutions"
        ]
        self.table_keywords = ["総計・自己合計・委託合計", "委託内訳", "法人内訳", "金融機関内訳"]
        self.table_names = ["Table1_Main_Summary", "Table2_Brokerage_Breakdown", "Table3_Institutions_Breakdown", "Table4_Financial_Breakdown"]
        
        logger.info(f"🚀 Selective Page Converter configured:")
        logger.info(f"  • Max workers: {max_workers}")
        logger.info(f"  • CPU cores available: {multiprocessing.cpu_count()}")
        logger.info(f"  • Will only convert pages containing: {self.target_subtitles}")
        logger.info(f"  • Input folder: {self.input_folder}")
        logger.info(f"  • Output folder: {self.output_folder}")

    def extract_table_titles_from_text(self, text: str) -> List[str]:
        """Extract table titles from page text."""
        found_titles = []
        for keyword, full_title in zip(self.table_keywords, self.table_section_titles):
            if keyword in text:
                found_titles.append(full_title)
        # Ensure we have titles for all expected tables
        while len(found_titles) < len(self.table_keywords):
            found_titles.append(f"Table Title {len(found_titles) + 1} (Not Found)")
        return found_titles

    def extract_table_titles_for_page(self, pdf_path: Path, page_number: int) -> List[str]:
        """Extract table titles for a specific page."""
        try:
            doc = fitz.open(str(pdf_path))
            page = doc[page_number]
            text = page.get_text("text")
            doc.close()
            return self.extract_table_titles_from_text(text)
        except Exception as e:
            logger.error(f"Could not extract table titles from page {page_number + 1}: {e}")
            return self.table_section_titles

    def get_relevant_pages_and_subtitles(self, pdf_path: Path) -> Dict[int, str]:
        """Scans a PDF and captures the FULL LINE of text for the subtitle."""
        relevant_pages = {}
        try:
            doc = fitz.open(str(pdf_path))
            logger.info(f"🔎 Scanning {pdf_path.name} ({len(doc)} pages) for relevant subtitles...")
            
            for i, page in enumerate(doc):
                page_match_found = False
                lines = page.get_text("text").splitlines()
                
                for line in lines:
                    if not line.strip():
                        continue

                    # Step 1: Check if this specific line is an excluded type
                    is_excluded_line = any(ex_keyword in line for ex_keyword in self.primary_exclusion_keywords)
                    if is_excluded_line:
                        continue
                        
                    # Step 2: If not excluded, check for a target subtitle
                    for subtitle_keyword in self.target_subtitles:
                        if subtitle_keyword in line:
                            # SUCCESS! Store the entire line instead of just the keyword
                            full_subtitle = line.strip()
                            logger.info(f"  > Page {i + 1}: MATCH for '{subtitle_keyword}' -> Capturing full title: '{full_subtitle}'")
                            relevant_pages[i] = full_subtitle
                            page_match_found = True
                            break # Found the best match for this page, move on
                    
                    if page_match_found:
                        break # Move to the next page
            doc.close()
        except Exception as e:
            logger.error(f"Could not scan PDF {pdf_path.name}: {e}")
        return relevant_pages

    def process_all_files(self):
        """Process all PDF files by selectively converting relevant pages."""
        if not self.input_folder.exists():
            logger.error(f"Input folder '{self.input_folder}' does not exist!")
            return
            
        pdf_files = list(self.input_folder.glob("*.pdf"))
        if not pdf_files:
            logger.error(f"No PDF files found in '{self.input_folder}'")
            return
            
        print(f"📁 Found {len(pdf_files)} PDF files to process.")
        
        self.process_pdfs_selectively(pdf_files)
        self.cleanup_temp_files()

    def process_pdfs_selectively(self, pdf_files: List[Path]):
        """Identifies relevant pages in PDFs and processes them in parallel."""
        start_time = time.time()
        all_page_jobs = []
        
        for pdf_file in pdf_files:
            relevant_pages = self.get_relevant_pages_and_subtitles(pdf_file)
            if not relevant_pages:
                logger.warning(f"⚠️ No relevant pages found in {pdf_file.name}. Skipping.")
                continue
            
            pdf_name = pdf_file.stem
            for page_num, subtitle in relevant_pages.items():
                # Extract table titles for this specific page
                table_titles = self.extract_table_titles_for_page(pdf_file, page_num)
                job = PageJob(
                    pdf_path=pdf_file,
                    pdf_name=pdf_name,
                    page_number=page_num,
                    subtitle=subtitle,  # This now contains the full subtitle line
                    job_id=f"{pdf_name}-P{page_num+1}",
                    chunks_folder=self.chunks_folder,
                    table_titles=table_titles
                )
                all_page_jobs.append(job)
        
        if not all_page_jobs:
            logger.info("No relevant pages to process across all files.")
            return

        total_pages_to_process = len(all_page_jobs)
        logger.info(f"⚡ Starting parallel conversion of {total_pages_to_process} relevant pages...")
        
        page_results = []
        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_job = {executor.submit(convert_page_worker, job): job for job in all_page_jobs}
            for future in as_completed(future_to_job):
                result = future.result()
                page_results.append(result)

        parallel_time = time.time() - start_time
        logger.info(f"⚡ Parallel conversion finished in {parallel_time:.2f} seconds.")

        # Group results by PDF
        pdf_results = {}
        for result in page_results:
            if result['success']:
                pdf_name = result['pdf_name']
                if pdf_name not in pdf_results:
                    pdf_results[pdf_name] = []
                pdf_results[pdf_name].append(result)
        
        # Combine pages and convert to Excel for each PDF
        for pdf_name, results in pdf_results.items():
            pdf_start_time = time.time()
            logger.info(f"📊 Processing results for {pdf_name}...")
            
            results.sort(key=lambda x: x['page_number'])
            
            page_paths = [Path(r['page_path']) for r in results]
            combined_docx_path = self.docx_folder / f"{pdf_name}.docx"
            
            if self.combine_pages(page_paths, combined_docx_path):
                self.convert_docx_to_excel(combined_docx_path, results)
                
            pdf_time = time.time() - pdf_start_time
            print(f"\n🎉 {pdf_name} COMPLETED in {pdf_time:.2f}s")
            print(f"  - Converted {len(results)} relevant pages.")
            print(f"  - Excel output: {self.output_folder / (pdf_name + '_extracted.xlsx')}")

    def combine_pages(self, page_paths: List[Path], output_path: Path) -> bool:
        """Combines multiple single-page DOCX files into one."""
        if not page_paths:
            return False
        
        try:
            # The first page becomes the base document
            combined_doc = Document(str(page_paths[0]))
            
            # Append the content of subsequent pages
            for path in page_paths[1:]:
                sub_doc = Document(str(path))
                for element in sub_doc.element.body:
                    combined_doc.element.body.append(element)
            
            combined_doc.save(str(output_path))
            logger.info(f"🔗 Combined {len(page_paths)} page(s) into {output_path.name}")
            return True
        except Exception as e:
            logger.error(f"Failed to combine pages into {output_path.name}: {e}")
            return False

    def convert_docx_to_excel(self, docx_path: Path, page_results: List[Dict]):
        """Converts the combined DOCX to a single Excel file with enhanced titles."""
        try:
            doc = Document(docx_path)
            total_tables = len(doc.tables)
            if total_tables == 0:
                logger.warning(f"⚠️ No tables found in {docx_path.name}")
                return

            # Create mapping from table index to subtitle and table titles
            table_to_subtitle_map = {}
            table_to_table_title_map = {}
            current_table_index = 0
            
            for result in page_results:
                num_tables_on_page = result['table_count']
                page_table_titles = result.get('table_titles', self.table_section_titles)
                
                for i in range(num_tables_on_page):
                    table_to_subtitle_map[current_table_index] = result['subtitle']  # Full subtitle
                    table_pos_on_page = i % 4
                    if table_pos_on_page < len(page_table_titles):
                        table_to_table_title_map[current_table_index] = page_table_titles[table_pos_on_page]
                    else:
                        table_to_table_title_map[current_table_index] = f"Table Title {table_pos_on_page + 1}"
                    current_table_index += 1
            
            wb = Workbook()
            wb.remove(wb.active)
            self.create_summary_sheet(wb, docx_path.name, len(page_results), total_tables)
            
            for i, table in enumerate(doc.tables):
                page_number_guess = (i // 4) + 1
                table_position = i % 4
                table_name = self.table_names[table_position]
                
                sheet_name = f"P{page_number_guess}_{table_name[:20]}"
                ws = wb.create_sheet(title=sheet_name)
                
                # Get the pre-identified subtitle and table title for this table
                subtitle = table_to_subtitle_map.get(i, "Subtitle Not Found")
                table_title = table_to_table_title_map.get(i, "Table Title Not Found")
                main_title = "Trading by Type of Investors"
                
                self.copy_table_to_sheet_with_enhanced_titles(table, ws, main_title, subtitle, table_title)
            
            # Save with original naming convention
            excel_path = self.output_folder / f"{docx_path.stem}_extracted.xlsx"
            wb.save(excel_path)
            
        except Exception as e:
            logger.error(f"❌ Error processing {docx_path.name} to Excel: {e}")

    def copy_table_to_sheet_with_enhanced_titles(self, table, ws, main_title, subtitle, table_title):
        """Copies a table to an Excel sheet with enhanced title information."""
        # Add enhanced title information
        ws['A1'] = f"Title: {main_title}"
        ws['A2'] = f"Subtitle: {subtitle}"        # Full captured subtitle
        ws['A3'] = f"Table Title: {table_title}"  # Extracted table title
        ws['A4'] = ""  # Spacer row
        
        start_row = 5
        for r, table_row in enumerate(table.rows):
            for c, cell in enumerate(table_row.cells):
                ws.cell(row=start_row + r, column=c + 1, value=cell.text.strip())
        
        self.apply_enhanced_formatting(ws, start_row, len(table.rows))

    def apply_enhanced_formatting(self, ws, start_row, num_rows):
        """Applies enhanced formatting to the Excel sheet."""
        title_fill = PatternFill(start_color='E6F3FF', end_color='E6F3FF', fill_type='solid')
        title_font = Font(bold=True, size=12)
        
        ws['A1'].font = title_font
        ws['A1'].fill = title_fill
        ws['A2'].font = Font(bold=True)
        ws['A3'].font = Font(bold=True)
        
        if num_rows > 0:
            header_fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
            for cell in ws[start_row]:
                if cell.value:
                    cell.font = Font(bold=True)
                    cell.fill = header_fill

    def create_summary_sheet(self, wb, filename, pages_converted, total_tables):
        """Creates a summary sheet with processing information."""
        ws = wb.create_sheet(title="Summary", index=0)
        ws['A1'] = "📋 Selective PDF→DOCX→Excel Conversion Summary"
        ws['A1'].font = Font(size=16, bold=True)
        ws['A3'] = f"Source File: {filename}"
        ws['A4'] = f"Pages Converted: {pages_converted}"
        ws['A5'] = f"Total Tables Extracted: {total_tables}"
        ws['A6'] = f"Processing Method: Selective Page Conversion with Enhanced Titles"

    def cleanup_temp_files(self):
        """Cleans up temporary page files."""
        try:
            chunk_files = list(self.chunks_folder.glob("*.docx"))
            for chunk_file in chunk_files:
                chunk_file.unlink()
            logger.info(f"🧹 Cleaned up {len(chunk_files)} temporary page files.")
        except Exception as e:
            logger.warning(f"Could not clean up temp files: {e}")


def main():
    print("📋 Title-Enhanced Selective Page PDF→DOCX→Excel Converter")
    print("=" * 75)
    print("🚀 This script will pre-scan PDFs and only convert pages with specific subtitles.")
    print("📁 Uses original folder structure: input/ → extracted_data/")
    print("=" * 75)
    
    max_workers = 6
    try:
        user_workers = input(f"Max workers (default {max_workers}): ").strip()
        if user_workers:
            max_workers = int(user_workers)
    except ValueError:
        print(f"Invalid input. Using default: {max_workers} workers.")
    
    # Use original folder structure
    converter = TitleEnhancedConverter(
        input_folder="input",
        output_folder="extracted_data", 
        max_workers=max_workers
    )
    converter.process_all_files()
    
    print(f"\n📁 Intermediate DOCX files saved in: {converter.docx_folder}")
    print(f"💾 Final Excel files saved in: {converter.output_folder}")

if __name__ == "__main__":
    main()